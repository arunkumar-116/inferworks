# import os
# import pickle
# from typing import List, Dict, Any, Tuple
# import numpy as np
# from sklearn.metrics.pairwise import cosine_similarity
# from models.embeddings import AzureEmbeddingModel
# import PyPDF2
# import docx
# import logging
# import re
# import csv
# import json
# import pandas as pd

# class DocumentProcessor:
#     def __init__(self):
#         self.embedding_model = AzureEmbeddingModel()
    
#     def extract_text_from_pdf(self, file_path: str) -> str:
#         """Extract text from PDF file"""
#         try:
#             text = ""
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text() + "\n"
#             return text
#         except Exception as e:
#             logging.error(f"Error extracting PDF text: {e}")
#             return ""
    
#     def extract_text_from_docx(self, file_path: str) -> str:
#         """Extract text from DOCX file"""
#         try:
#             doc = docx.Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#             return text
#         except Exception as e:
#             logging.error(f"Error extracting DOCX text: {e}")
#             return ""
    
#     def extract_text_from_txt(self, file_path: str) -> str:
#         """Extract text from TXT file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         except Exception as e:
#             logging.error(f"Error extracting TXT text: {e}")
#             return ""
    
#     def extract_text_from_csv(self, file_path: str) -> str:
#         """Extract text from CSV file with medical data context"""
#         try:
#             text = ""
#             # Try to read with pandas first for better handling
#             try:
#                 df = pd.read_csv(file_path)
#                 text = self._format_csv_data(df, file_path)
#             except:
#                 # Fallback to standard CSV reader
#                 with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
#                     csv_reader = csv.reader(file)
#                     headers = next(csv_reader, [])
#                     text += f"CSV File: {os.path.basename(file_path)}\n"
#                     text += f"Headers: {', '.join(headers)}\n\n"
                    
#                     for i, row in enumerate(csv_reader):
#                         if i < 10:  # Limit to first 10 rows for context
#                             text += f"Row {i+1}: {', '.join(str(cell) for cell in row)}\n"
            
#             return text
#         except Exception as e:
#             logging.error(f"Error extracting CSV text: {e}")
#             return ""
    
#     def _format_csv_data(self, df: pd.DataFrame, file_path: str) -> str:
#         """Format CSV data for better medical context"""
#         text = f"Medical Data from {os.path.basename(file_path)}:\n\n"
        
#         # Add basic info
#         text += f"Shape: {df.shape[0]} rows, {df.shape[1]} columns\n"
        
#         # Add column descriptions if available
#         text += "Columns:\n"
#         for col in df.columns:
#             text += f"- {col}: {df[col].dtype}, Sample: {df[col].iloc[0] if not df.empty else 'N/A'}\n"
        
#         text += "\nSample Data:\n"
#         # Include first few rows for context
#         for i in range(min(5, len(df))):
#             text += f"Row {i+1}: "
#             row_data = []
#             for col in df.columns:
#                 value = str(df[col].iloc[i]) if not pd.isna(df[col].iloc[i]) else "N/A"
#                 row_data.append(f"{col}={value}")
#             text += "; ".join(row_data) + "\n"
        
#         return text
    
#     def extract_text_from_json(self, file_path: str) -> str:
#         """Extract text from JSON file with medical data context"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 data = json.load(file)
            
#             text = f"JSON Data from {os.path.basename(file_path)}:\n\n"
            
#             if isinstance(data, list):
#                 # Handle array of objects (common in medical data)
#                 text += f"Contains {len(data)} records\n\n"
#                 for i, item in enumerate(data[:3]):  # First 3 items
#                     text += f"Record {i+1}:\n"
#                     text += self._format_json_item(item, indent=2)
#                     text += "\n"
#             elif isinstance(data, dict):
#                 # Handle single object
#                 text += self._format_json_item(data, indent=0)
#             else:
#                 text += str(data)
            
#             return text
#         except Exception as e:
#             logging.error(f"Error extracting JSON text: {e}")
#             return ""
    
#     def _format_json_item(self, item: Any, indent: int = 0) -> str:
#         """Recursively format JSON item for readable text"""
#         if isinstance(item, dict):
#             text = ""
#             for key, value in item.items():
#                 indent_str = " " * indent
#                 if isinstance(value, (dict, list)):
#                     text += f"{indent_str}{key}:\n"
#                     text += self._format_json_item(value, indent + 2)
#                 else:
#                     text += f"{indent_str}{key}: {value}\n"
#             return text
#         elif isinstance(item, list):
#             text = ""
#             for i, value in enumerate(item):
#                 indent_str = " " * indent
#                 if isinstance(value, (dict, list)):
#                     text += f"{indent_str}Item {i+1}:\n"
#                     text += self._format_json_item(value, indent + 2)
#                 else:
#                     text += f"{indent_str}Item {i+1}: {value}\n"
#             return text
#         else:
#             return " " * indent + str(item) + "\n"
    
#     def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
#         """Split text into chunks with overlap, preserving medical terminology context"""
#         # First, split by sections that might be important in medical documents
#         sections = re.split(r'\n\s*(?:ABSTRACT|INTRODUCTION|METHODS|RESULTS|DISCUSSION|CONCLUSION|REFERENCES)\s*\n', text, flags=re.IGNORECASE)
        
#         chunks = []
#         for section in sections:
#             if not section.strip():
#                 continue
                
#             # Further chunk the section if it's too long
#             if len(section) > chunk_size:
#                 # Try to split at sentence boundaries for better context preservation
#                 sentences = re.split(r'(?<=[.!?])\s+', section)
#                 current_chunk = ""
                
#                 for sentence in sentences:
#                     if len(current_chunk) + len(sentence) < chunk_size:
#                         current_chunk += sentence + " "
#                     else:
#                         if current_chunk:
#                             chunks.append(current_chunk.strip())
#                         current_chunk = sentence + " "
                
#                 if current_chunk:
#                     chunks.append(current_chunk.strip())
#             else:
#                 chunks.append(section.strip())
        
#         return chunks

# class RAGSystem:
#     def __init__(self):
#         self.document_processor = DocumentProcessor()
#         self.vector_store = VectorStore()
    
#     def process_and_store_documents(self, file_paths: List[str]):
#         """Process and store documents in vector store"""
#         documents = []
        
#         for file_path in file_paths:
#             if not os.path.exists(file_path):
#                 continue
            
#             # Extract text based on file type
#             file_ext = os.path.splitext(file_path)[1].lower()
#             text = ""
            
#             if file_ext == '.pdf':
#                 text = self.document_processor.extract_text_from_pdf(file_path)
#             elif file_ext == '.docx':
#                 text = self.document_processor.extract_text_from_docx(file_path)
#             elif file_ext == '.txt':
#                 text = self.document_processor.extract_text_from_txt(file_path)
#             elif file_ext == '.csv':
#                 text = self.document_processor.extract_text_from_csv(file_path)
#             elif file_ext == '.json':
#                 text = self.document_processor.extract_text_from_json(file_path)
#             else:
#                 logging.warning(f"Unsupported file format: {file_ext}")
#                 continue
            
#             if not text.strip():
#                 logging.warning(f"No text extracted from: {file_path}")
#                 continue
            
#             # Chunk the text with medical context preservation
#             chunks = self.document_processor.chunk_text(text)
            
#             # Create document objects
#             for i, chunk in enumerate(chunks):
#                 documents.append({
#                     'content': chunk,
#                     'source': file_path,
#                     'chunk_id': i,
#                     'metadata': {
#                         'file_type': file_ext,
#                         'is_medical_content': self._is_medical_content(chunk)
#                     }
#                 })
        
#         # Add to vector store
#         if documents:
#             self.vector_store.add_documents(documents)
#         return len(documents)
    
#     def _is_medical_content(self, text: str) -> bool:
#         """Simple check to determine if content appears to be medical"""
#         medical_indicators = [
#             'patient', 'treatment', 'diagnosis', 'symptom', 'disease', 
#             'clinical', 'therapy', 'medication', 'dosage', 'mg/kg',
#             'methodology', 'results', 'conclusion', 'abstract', 'study',
#             'medical', 'health', 'hospital', 'doctor', 'nurse', 'physician',
#             'blood', 'pressure', 'heart', 'rate', 'temperature', 'respiratory',
#             'spo2', 'diagnosis', 'medications', 'admission', 'discharge'
#         ]
        
#         text_lower = text.lower()
#         indicator_count = sum(1 for indicator in medical_indicators if indicator in text_lower)
        
#         return indicator_count >= 2  # Lower threshold for CSV/JSON data
    
#     def retrieve_relevant_context(self, query: str, top_k: int = 3) -> Tuple[str, List[str]]:
#         """Retrieve relevant context for RAG and return context with unique source info"""
#         results = self.vector_store.search(query, top_k)
        
#         if not results:
#             return "", []
        
#         context = "Relevant information from research documents:\n\n"
#         source_files = {}  # Use dict to track files with their highest similarity
        
#         for i, result in enumerate(results, 1):
#             file_name = os.path.basename(result['source'])
#             context += f"Context {i} (from {file_name}):\n"
#             context += f"{result['content']}\n\n"
            
#             # Track the highest similarity score for each file
#             if file_name not in source_files or result['similarity'] > source_files[file_name]:
#                 source_files[file_name] = result['similarity']
        
#         # Sort sources by similarity score (highest first) and return just the file names
#         sorted_sources = sorted(source_files.keys(), key=lambda x: source_files[x], reverse=True)
#         return context, sorted_sources

import os
import pickle
from typing import List, Dict, Any, Tuple
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from models.embeddings import AzureEmbeddingModel
import PyPDF2
import docx
import logging
import re
import csv
import json
import pandas as pd
from io import StringIO

class DocumentProcessor:
    def __init__(self):
        self.embedding_model = AzureEmbeddingModel()
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logging.error(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logging.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logging.error(f"Error extracting TXT text: {e}")
            return ""
    
    def extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file with medical data context - reads entire dataset"""
        try:
            text = ""
            
            # Read the entire CSV file
            try:
                # Try with pandas first for better handling of large files
                df = pd.read_csv(file_path, low_memory=False)
                text = self._format_complete_csv_data(df, file_path)
            except Exception as pd_error:
                logging.warning(f"Pandas read failed, falling back to CSV reader: {pd_error}")
                # Fallback to standard CSV reader for the entire file
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    csv_reader = csv.reader(file)
                    headers = next(csv_reader, [])
                    text += f"CSV File: {os.path.basename(file_path)}\n"
                    text += f"Headers: {', '.join(headers)}\n\n"
                    text += f"Total rows: {sum(1 for row in csv_reader) + 1}\n\n"
                    
                    # Reset and read all data
                    file.seek(0)
                    next(csv_reader)  # Skip headers again
                    
                    all_rows = []
                    for row in csv_reader:
                        if row:  # Skip empty rows
                            all_rows.append(row)
                    
                    # Include statistical summary for large files
                    if len(all_rows) > 50:
                        text += f"Dataset too large to display all {len(all_rows)} rows. Showing statistical summary:\n\n"
                        text += self._generate_csv_summary(headers, all_rows)
                    else:
                        text += "Complete dataset:\n\n"
                        for i, row in enumerate(all_rows):
                            text += f"Row {i+1}: {', '.join(str(cell) for cell in row)}\n"
            
            return text
        except Exception as e:
            logging.error(f"Error extracting CSV text: {e}")
            return ""
    
    def _format_complete_csv_data(self, df: pd.DataFrame, file_path: str) -> str:
        """Format complete CSV data for medical context"""
        text = f"Complete Medical Dataset from {os.path.basename(file_path)}:\n\n"
        
        # Basic dataset info
        text += f"Dataset Shape: {df.shape[0]} rows, {df.shape[1]} columns\n"
        
        # Column information with data types and sample values
        text += "Column Details:\n"
        for col in df.columns:
            non_null_count = df[col].count()
            null_count = df.shape[0] - non_null_count
            dtype = df[col].dtype
            sample_value = df[col].iloc[0] if non_null_count > 0 else "N/A"
            
            text += f"- {col}: {dtype}, Non-null: {non_null_count}, Null: {null_count}, Sample: {sample_value}\n"
        
        text += "\n"
        
        # Statistical summary for numerical columns
        numerical_cols = df.select_dtypes(include=[np.number]).columns
        if len(numerical_cols) > 0:
            text += "Numerical Column Statistics:\n"
            for col in numerical_cols:
                stats = df[col].describe()
                text += f"  {col}: mean={stats.get('mean', 'N/A'):.2f}, min={stats.get('min', 'N/A')}, max={stats.get('max', 'N/A')}\n"
            text += "\n"
        
        # For large datasets, create meaningful chunks instead of showing all data
        if len(df) > 100:
            text += f"Large dataset with {len(df)} rows. Creating analytical chunks:\n\n"
            
            # Create chunks based on meaningful groupings if possible
            if 'Medical Condition' in df.columns:
                text += "Summary by Medical Condition:\n"
                condition_counts = df['Medical Condition'].value_counts()
                for condition, count in condition_counts.items():
                    text += f"  {condition}: {count} patients\n"
                text += "\n"
            
            if 'Age' in df.columns and pd.api.types.is_numeric_dtype(df['Age']):
                text += "Age Distribution:\n"
                age_stats = df['Age'].describe()
                text += f"  Mean: {age_stats['mean']:.1f}, Range: {age_stats['min']}-{age_stats['max']}\n"
                text += "\n"
            
            # Include sample of data from different parts of the dataset
            text += "Data Samples from different segments:\n"
            sample_size = min(10, len(df) // 10)
            for i in range(0, len(df), len(df) // sample_size):
                if i < len(df):
                    sample_row = df.iloc[i]
                    text += f"Row {i+1}: {', '.join([f'{col}={sample_row[col]}' for col in df.columns[:3]])}...\n"
        
        else:
            # For smaller datasets, include all data in structured format
            text += "Complete Dataset:\n\n"
            for i, (_, row) in enumerate(df.iterrows()):
                row_data = []
                for col in df.columns:
                    value = str(row[col]) if not pd.isna(row[col]) else "N/A"
                    # Truncate very long values
                    if len(value) > 100:
                        value = value[:100] + "..."
                    row_data.append(f"{col}={value}")
                text += f"Row {i+1}: {' | '.join(row_data)}\n"
        
        return text
    
    def _generate_csv_summary(self, headers: List[str], all_rows: List[List[str]]) -> str:
        """Generate statistical summary for large CSV files"""
        text = "Statistical Summary:\n\n"
        
        # Convert to DataFrame for easier analysis
        try:
            df = pd.DataFrame(all_rows, columns=headers)
            
            # Basic stats
            text += f"Total records: {len(df)}\n"
            text += f"Columns: {len(headers)}\n\n"
            
            # Column analysis
            text += "Column Analysis:\n"
            for col in headers:
                unique_count = df[col].nunique()
                null_count = df[col].isnull().sum()
                sample_value = df[col].iloc[0] if len(df) > 0 else "N/A"
                
                text += f"  {col}: {unique_count} unique values, {null_count} nulls, sample: {sample_value}\n"
            
            # Identify potential medical columns
            medical_columns = []
            for col in headers:
                col_lower = col.lower()
                if any(term in col_lower for term in ['medical', 'patient', 'diagnosis', 'treatment', 'medication', 'symptom']):
                    medical_columns.append(col)
            
            if medical_columns:
                text += f"\nKey Medical Columns: {', '.join(medical_columns)}\n"
                
                # Summary for key medical columns
                for col in medical_columns[:3]:  # Limit to top 3
                    if df[col].nunique() < 20:  # For categorical data
                        value_counts = df[col].value_counts().head(5)
                        text += f"  {col} distribution: {dict(value_counts)}\n"
            
        except Exception as e:
            logging.warning(f"Could not generate detailed summary: {e}")
            text += f"Raw data contains {len(all_rows)} rows with {len(headers)} columns\n"
        
        return text
    
    def extract_text_from_json(self, file_path: str) -> str:
        """Extract text from JSON file with medical data context - reads entire dataset"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            text = f"Complete JSON Data from {os.path.basename(file_path)}:\n\n"
            
            if isinstance(data, list):
                # Handle array of objects (common in medical data)
                text += f"Contains {len(data)} complete records\n\n"
                
                if len(data) > 50:
                    text += f"Large dataset with {len(data)} records. Creating analytical summary:\n\n"
                    text += self._analyze_json_array(data)
                else:
                    # Include all records for smaller datasets
                    for i, item in enumerate(data):
                        text += f"Record {i+1}:\n"
                        text += self._format_json_item(item, indent=2)
                        text += "\n"
                        
            elif isinstance(data, dict):
                # Handle single object or nested structure
                if any(isinstance(v, (list, dict)) for v in data.values()):
                    text += "Complex nested structure analyzed:\n\n"
                    text += self._analyze_complex_json(data)
                else:
                    text += "Complete object data:\n\n"
                    text += self._format_json_item(data, indent=0)
            else:
                text += f"Complete data: {str(data)}"
            
            return text
        except Exception as e:
            logging.error(f"Error extracting JSON text: {e}")
            return ""
    
    def _analyze_json_array(self, data: List[Any]) -> str:
        """Analyze large JSON arrays for medical context"""
        text = ""
        
        if not data or not isinstance(data[0], dict):
            return f"Array of {len(data)} simple elements\n"
        
        # Analyze structure of first few items
        sample_item = data[0]
        text += f"Record structure: {list(sample_item.keys())}\n\n"
        
        # Count records by key medical fields if present
        medical_fields = ['diagnosis', 'condition', 'treatment', 'medication', 'status']
        found_medical_fields = [field for field in medical_fields if field in sample_item]
        
        if found_medical_fields:
            text += "Medical Field Analysis:\n"
            for field in found_medical_fields[:3]:  # Limit to top 3
                try:
                    values = [str(item.get(field, 'N/A')) for item in data if field in item]
                    if values:
                        unique_values = set(values)
                        if len(unique_values) <= 20:
                            text += f"  {field}: {len(unique_values)} unique values\n"
                            # Show top 5 values
                            from collections import Counter
                            counter = Counter(values)
                            for value, count in counter.most_common(5):
                                text += f"    - {value}: {count} records\n"
                        else:
                            text += f"  {field}: {len(unique_values)} unique values (too many to display)\n"
                except:
                    pass
            text += "\n"
        
        # Include sample records from different parts of the dataset
        text += "Sample Records:\n"
        sample_indices = [0, len(data)//4, len(data)//2, 3*len(data)//4, -1]
        for idx in sample_indices:
            if 0 <= idx < len(data):
                item = data[idx]
                text += f"Record {idx+1}:\n"
                # Show only key fields for brevity
                key_fields = list(item.keys())[:5]  # First 5 fields
                for field in key_fields:
                    value = str(item.get(field, 'N/A'))
                    if len(value) > 100:
                        value = value[:100] + "..."
                    text += f"  {field}: {value}\n"
                text += "\n"
        
        return text
    
    def _analyze_complex_json(self, data: Dict[str, Any]) -> str:
        """Analyze complex nested JSON structures"""
        text = ""
        
        def analyze_node(node, path="", depth=0):
            nonlocal text
            indent = "  " * depth
            
            if isinstance(node, dict):
                text += f"{indent}{path}: Object with keys {list(node.keys())}\n"
                for key, value in node.items():
                    new_path = f"{path}.{key}" if path else key
                    analyze_node(value, new_path, depth + 1)
            elif isinstance(node, list):
                text += f"{indent}{path}: Array with {len(node)} elements\n"
                if node and depth < 3:  # Limit recursion depth
                    analyze_node(node[0], f"{path}[0]", depth + 1)
            else:
                sample = str(node)
                if len(sample) > 50:
                    sample = sample[:50] + "..."
                text += f"{indent}{path}: {type(node).__name__} = {sample}\n"
        
        analyze_node(data)
        return text
    
    def _format_json_item(self, item: Any, indent: int = 0) -> str:
        """Recursively format JSON item for readable text"""
        if isinstance(item, dict):
            text = ""
            for key, value in item.items():
                indent_str = " " * indent
                if isinstance(value, (dict, list)):
                    text += f"{indent_str}{key}:\n"
                    text += self._format_json_item(value, indent + 2)
                else:
                    value_str = str(value)
                    if len(value_str) > 200:
                        value_str = value_str[:200] + "..."
                    text += f"{indent_str}{key}: {value_str}\n"
            return text
        elif isinstance(item, list):
            text = ""
            for i, value in enumerate(item):
                indent_str = " " * indent
                if isinstance(value, (dict, list)) and i < 5:  # Limit nested items
                    text += f"{indent_str}Item {i+1}:\n"
                    text += self._format_json_item(value, indent + 2)
                elif i < 10:  # Limit total items shown
                    value_str = str(value)
                    if len(value_str) > 100:
                        value_str = value_str[:100] + "..."
                    text += f"{indent_str}Item {i+1}: {value_str}\n"
                elif i == 10:
                    text += f"{indent_str}... and {len(item) - 10} more items\n"
            return text
        else:
            value_str = str(item)
            if len(value_str) > 200:
                value_str = value_str[:200] + "..."
            return " " * indent + value_str + "\n"
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into chunks with overlap, preserving medical terminology context"""
        # For very large texts, use more sophisticated chunking
        if len(text) > 10000:
            return self._chunk_large_text(text, chunk_size, overlap)
        
        # First, split by sections that might be important in medical documents
        sections = re.split(r'\n\s*(?:ABSTRACT|INTRODUCTION|METHODS|RESULTS|DISCUSSION|CONCLUSION|REFERENCES)\s*\n', text, flags=re.IGNORECASE)
        
        chunks = []
        for section in sections:
            if not section.strip():
                continue
                
            # Further chunk the section if it's too long
            if len(section) > chunk_size:
                # Try to split at sentence boundaries for better context preservation
                sentences = re.split(r'(?<=[.!?])\s+', section)
                current_chunk = ""
                
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) < chunk_size:
                        current_chunk += sentence + " "
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = sentence + " "
                
                if current_chunk:
                    chunks.append(current_chunk.strip())
            else:
                chunks.append(section.strip())
        
        return chunks
    
    def _chunk_large_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """Chunk very large texts efficiently"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk end
                sentence_end = text.rfind('.', start, end)
                paragraph_end = text.rfind('\n\n', start, end)
                
                if paragraph_end > start and paragraph_end - start > chunk_size // 2:
                    end = paragraph_end + 2
                elif sentence_end > start and sentence_end - start > chunk_size // 2:
                    end = sentence_end + 1
                # Otherwise, look for any reasonable break point
                else:
                    space_break = text.rfind(' ', start, end)
                    if space_break > start and space_break - start > chunk_size // 2:
                        end = space_break + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position, considering overlap
            start = end - overlap if end - overlap > start else end
        
        return chunks

class VectorStore:
    def __init__(self, store_path: str = "./vector_store"):
        self.store_path = store_path
        self.embedding_model = AzureEmbeddingModel()
        self.documents = []
        self.embeddings = []
        self.load_store()
    
    def load_store(self):
        """Load existing vector store"""
        try:
            if os.path.exists(f"{self.store_path}/documents.pkl"):
                with open(f"{self.store_path}/documents.pkl", 'rb') as f:
                    self.documents = pickle.load(f)
                with open(f"{self.store_path}/embeddings.pkl", 'rb') as f:
                    self.embeddings = pickle.load(f)
        except Exception as e:
            logging.error(f"Error loading vector store: {e}")
    
    def save_store(self):
        """Save vector store"""
        try:
            os.makedirs(self.store_path, exist_ok=True)
            with open(f"{self.store_path}/documents.pkl", 'wb') as f:
                pickle.dump(self.documents, f)
            with open(f"{self.store_path}/embeddings.pkl", 'wb') as f:
                pickle.dump(self.embeddings, f)
        except Exception as e:
            logging.error(f"Error saving vector store: {e}")
    
    def add_documents(self, documents: List[Dict[str, Any]]):
        """Add documents to vector store"""
        texts = [doc['content'] for doc in documents]
        embeddings = self.embedding_model.get_embeddings(texts)
        
        for doc, embedding in zip(documents, embeddings):
            self.documents.append(doc)
            self.embeddings.append(embedding)
        
        self.save_store()
    
    def search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar documents with medical context boosting"""
        if not self.documents:
            return []
        
        query_embedding = self.embedding_model.get_single_embedding(query)
        if not query_embedding:
            return []
        
        similarities = cosine_similarity([query_embedding], self.embeddings)[0]
        
        # Boost similarity for documents containing medical terminology
        medical_terms = self._extract_medical_terms(query)
        if medical_terms:
            for i, doc in enumerate(self.documents):
                doc_text = doc['content'].lower()
                term_count = sum(1 for term in medical_terms if term.lower() in doc_text)
                if term_count > 0:
                    similarities[i] *= (1 + 0.1 * term_count)  # Boost by 10% per term
        
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        results = []
        for idx in top_indices:
            results.append({
                **self.documents[idx],
                'similarity': similarities[idx]
            })
        
        return results
    
    def _extract_medical_terms(self, text: str) -> List[str]:
        """Simple medical term extractor - can be enhanced with a proper medical dictionary"""
        # Common medical prefixes, suffixes, and patterns
        medical_patterns = [
            r'\b[a-z]+ology\b', r'\b[a-z]+itis\b', r'\b[a-z]+ectomy\b', 
            r'\b[a-z]+pathy\b', r'\b[a-z]+emia\b', r'\b[a-z]+osis\b',
            r'\b[a-z]+matic\b', r'\b[a-z]+genic\b'
        ]
        
        terms = set()
        for pattern in medical_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            terms.update(matches)
        
        return list(terms)

class RAGSystem:
    def __init__(self):
        self.document_processor = DocumentProcessor()
        self.vector_store = VectorStore()
    
    def process_and_store_documents(self, file_paths: List[str]):
        """Process and store documents in vector store"""
        documents = []
        
        for file_path in file_paths:
            if not os.path.exists(file_path):
                continue
            
            # Extract text based on file type
            file_ext = os.path.splitext(file_path)[1].lower()
            text = ""
            
            if file_ext == '.pdf':
                text = self.document_processor.extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                text = self.document_processor.extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                text = self.document_processor.extract_text_from_txt(file_path)
            elif file_ext == '.csv':
                text = self.document_processor.extract_text_from_csv(file_path)
            elif file_ext == '.json':
                text = self.document_processor.extract_text_from_json(file_path)
            else:
                logging.warning(f"Unsupported file format: {file_ext}")
                continue
            
            if not text.strip():
                logging.warning(f"No text extracted from: {file_path}")
                continue
            
            # Chunk the text with medical context preservation
            chunks = self.document_processor.chunk_text(text)
            
            # Create document objects
            for i, chunk in enumerate(chunks):
                documents.append({
                    'content': chunk,
                    'source': file_path,
                    'chunk_id': i,
                    'metadata': {
                        'file_type': file_ext,
                        'is_medical_content': self._is_medical_content(chunk)
                    }
                })
        
        # Add to vector store
        if documents:
            self.vector_store.add_documents(documents)
        return len(documents)
    
    def _is_medical_content(self, text: str) -> bool:
        """Simple check to determine if content appears to be medical"""
        medical_indicators = [
            'patient', 'treatment', 'diagnosis', 'symptom', 'disease', 
            'clinical', 'therapy', 'medication', 'dosage', 'mg/kg',
            'methodology', 'results', 'conclusion', 'abstract', 'study',
            'medical', 'health', 'hospital', 'doctor', 'nurse', 'physician',
            'blood', 'pressure', 'heart', 'rate', 'temperature', 'respiratory',
            'spo2', 'diagnosis', 'medications', 'admission', 'discharge',
            'insurance', 'billing', 'room', 'condition', 'record', 'data'
        ]
        
        text_lower = text.lower()
        indicator_count = sum(1 for indicator in medical_indicators if indicator in text_lower)
        
        return indicator_count >= 2
    
    def retrieve_relevant_context(self, query: str, top_k: int = 3) -> Tuple[str, List[str]]:
        """Retrieve relevant context for RAG and return context with unique source info"""
        results = self.vector_store.search(query, top_k)
        
        if not results:
            return "", []
        
        context = "Relevant information from research documents:\n\n"
        source_files = {}  # Use dict to track files with their highest similarity
        
        for i, result in enumerate(results, 1):
            file_name = os.path.basename(result['source'])
            context += f"Context {i} (from {file_name}):\n"
            context += f"{result['content']}\n\n"
            
            # Track the highest similarity score for each file
            if file_name not in source_files or result['similarity'] > source_files[file_name]:
                source_files[file_name] = result['similarity']
        
        # Sort sources by similarity score (highest first) and return just the file names
        sorted_sources = sorted(source_files.keys(), key=lambda x: source_files[x], reverse=True)
        return context, sorted_sources